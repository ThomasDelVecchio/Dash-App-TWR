import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import plotly.graph_objects as go
import dash_wrappers as dw 
from tax_engine import build_tax_lots

# =====================================================================
# Formatting Helpers
# =====================================================================

# Global Asset Class Abbreviation Map
ASSET_CLASS_SHORT_MAP = {
    "US Large Cap": "US LC",
    "US Growth": "US Growth", 
    "US Small Cap": "US SC",
    "International Equity": "INTL EQTY",
    "Gold / Precious Metals": "GOLD",
    "Digital Assets": "DIGITAL",
    "US Bonds": "US Bonds", 
    "CASH": "CASH", 
    "Fixed Income": "FI"
}

def fmt_pct_clean(x):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"{float(x)*100:.2f}%"
    except:
        return "N/A"

def fmt_dollar_clean(x):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"${float(x):,.2f}"
    except:
        return "N/A"

def fmt_number_clean(x):
    try:
        if x is None or pd.isna(x):
            return "N/A"
        return f"{float(x):,.2f}"
    except:
        return "N/A"

def safe(x):
    return "N/A" if x is None or pd.isna(x) else x

# =====================================================================
# Document Styling Helpers
# =====================================================================

def set_narrow_margins(doc):
    """Sets page margins to 0.3 inches."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

def add_header(doc, text, level=1):
    """Adds a centered header."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Keep with next paragraph (the table or chart)
    p.paragraph_format.keep_with_next = True

def add_paragraph_centered(doc, text, bold=False, keep_after=False):
    """Adds a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if keep_after:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if bold:
        run.bold = True

def add_page_break(doc):
    doc.add_page_break()

def add_markdown_paragraph(doc, text):
    """
    Parses simple markdown (bold only via **) and adds to doc.
    """
    if not text: return
    
    # Split into logical paragraphs by newlines
    paragraphs = text.split('\n')
    
    for paragraph_text in paragraphs:
        if not paragraph_text.strip():
            continue
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Split by ** for bolding
        parts = paragraph_text.split("**")
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1: # Odd parts are inside **...**
                run.bold = True

# =====================================================================
# Table Helpers
# =====================================================================

def add_table(doc, headers, rows, right_align=None):
    """
    Adds a table with 'Light Grid Accent 1' style, autofitted to ~7.5 inches.
    """
    try:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Enable autofit
        table.autofit = True
        table.allow_autofit = True

        # Target width = 7.5" (8.5 - 0.5 - 0.5)
        max_width = Inches(7.5)
        col_width = max_width / len(headers)

        for col in table.columns:
            for cell in col.cells:
                cell.width = col_width

        # Header
        hdr_row = table.rows[0]
        # Repeat header row on new pages
        trPr = hdr_row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:tblHeader"))
        
        hdr = hdr_row.cells
        for i, h in enumerate(headers):
            hdr[i].text = str(h)
            for p in hdr[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True

        # Rows
        for row in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                
                # Align
                align = WD_ALIGN_PARAGRAPH.LEFT
                if right_align and i in right_align:
                    align = WD_ALIGN_PARAGRAPH.RIGHT
                elif i == 0 and len(headers) > 1: # Usually left align 1st col
                    align = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                
                for p in row_cells[i].paragraphs:
                    p.alignment = align

        # Prevent row splitting
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cant = OxmlElement("w:cantSplit")
            trPr.append(cant)
            
        doc.add_paragraph() # Spacer
        return table
    except Exception as e:
        print(f"Error adding table: {e}")
        return None

# =====================================================================
# Image / Chart Helpers
# =====================================================================

def add_figure_to_doc(doc, fig, width_inches=7.5, height_inches=4.5):
    """
    Converts a Plotly fig to image and adds it to the Word doc.
    Uses kaleido.
    """
    if fig is None:
        return
    
    # 1. Update layout for static export (white background)
    fig.update_layout(
         template='plotly_white',
         paper_bgcolor='white',
         plot_bgcolor='white',
         font=dict(size=10, color='black'),
         margin=dict(l=20, r=20, t=40, b=20),
    )
    
    # 2. Convert to image bytes
    try:
        img_bytes = fig.to_image(format="png", width=width_inches*100, height=height_inches*100, scale=3)
        
        # 3. Add to doc
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(BytesIO(img_bytes), width=Inches(width_inches))
    except Exception as e:
        err_p = doc.add_paragraph()
        err_p.add_run(f"[Chart could not be generated: {str(e)}]")


# =====================================================================
# Main Generator
# =====================================================================

def generate_word_report(data, sections, report_title, subtitle, period_label, mobile_mode=False, start_date=None, end_date=None):
    """
    Generates a Word Document based on selected sections.
    Args:
        end_date (datetime): Time Machine cutoff for tax lot calculations.
        mobile_mode (bool): If True, forces 1 chart per page and full width for readability.
        start_date (datetime): Optional anchor for period-based AI summaries.
    """
    doc = Document()
    set_narrow_margins(doc)
    
    # Helper to manage layout changes
    def check_break():
        if mobile_mode:
            # Continuous scroll effect for mobile (Spacer instead of Break)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("─" * 20).bold = True # Horizontal rule
            doc.add_paragraph() # Extra spacer
        # Note: Desktop layout relies on explicit breaks or flow as designed previously.

    # Standard width is 7.5" (8.5 page - 1.0 margins). 
    # In split mode (desktop default), we use 6" for side-by-side feel or smaller figs.
    # In mobile mode, we use full 7.5" for everything.
    w_full = 7.5
    w_split = 7.5 if mobile_mode else 6.0
    h_std = 5.0 if mobile_mode else 4.5
    
    # Title Page / Header
    add_header(doc, report_title, level=1)
    if subtitle:
        add_paragraph_centered(doc, subtitle)
    add_paragraph_centered(doc, f"Period: {period_label}", bold=True)
    add_paragraph_centered(doc, f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph() # Spacer

    # =================================================================
    # Dynamic Section Loop
    # =================================================================
    for section_key in sections:
    
        # -----------------------------------------------------------------
        # 0. Morning AI Summary (Optional)
        # -----------------------------------------------------------------
        if section_key == "morning_brief":
            # Avoid circular import at top level
            try:
                 from components.ai_brief import generate_ai_summary, generate_ai_summary_period
            except:
                 pass 

            # If we have the period-aware generator, prefer it (used in Custom Reports)
            if 'generate_ai_summary_period' in locals():
                summary_text = generate_ai_summary_period(data, start_date=start_date, end_date=end_date)
                header_text = "Period Performance Summary"
            elif 'generate_ai_summary' in locals():
                summary_text = generate_ai_summary(data)
                header_text = "Morning AI Summary"
            else:
                 summary_text = "AI Summary not available."
                 header_text = "AI Summary"
            
            check_break()
            add_header(doc, header_text, level=2)
            add_markdown_paragraph(doc, summary_text)
            doc.add_paragraph() # Spacer

        # -----------------------------------------------------------------
        # 1. Executive Summary
        # -----------------------------------------------------------------
        elif section_key == "summary":
            check_break()
            add_header(doc, "Executive Summary", level=2)
            
            metrics = dw.get_snapshot_metrics(data)
            metrics2 = data.get("snapshot_metrics", {})
            
            # Calculate MTD return from twr_df if available
            mtd_ret = 0.0
            twr_df = data.get("twr_df", pd.DataFrame())
            if not twr_df.empty and "Horizon" in twr_df.columns:
                row = twr_df[twr_df["Horizon"] == "MTD"]
                if not row.empty:
                    mtd_ret = row["Return"].iloc[0]

            # Get Current Value (PV)
            pv = data.get("pv", pd.Series())
            curr_val = pv.iloc[-1] if not pv.empty else 0.0
            
            kpis = [
                ["Metric", "Value"],
                ["Total P/L (SI)", fmt_dollar_clean(metrics.get('pl_si'))],
                ["Cumulative Return (SI)", fmt_pct_clean(metrics.get('twr_si'))],
                ["Current Value", fmt_dollar_clean(curr_val)],
                ["MTD Return", fmt_pct_clean(mtd_ret)],
                ["Max Drawdown", fmt_pct_clean(metrics.get('max_dd'))],
                ["Sharpe Ratio", f"{metrics.get('sharpe', 'N/A'):.2f}" if isinstance(metrics.get('sharpe'), (int, float)) else str(metrics.get('sharpe', 'N/A'))],
                ["Sortino Ratio", f"{metrics.get('sortino', 'N/A'):.2f}" if isinstance(metrics.get('sortino'), (int, float)) else str(metrics.get('sortino', 'N/A'))],
            ]
            add_table(doc, kpis[0], kpis[1:], right_align=[1])

        # -----------------------------------------------------------------
        # 2. Performance Chart
        # -----------------------------------------------------------------
        elif section_key == "performance_chart":
            check_break()
            add_header(doc, "Portfolio Performance", level=2)
            fig = dw.get_pv_mountain_chart(data, theme='light') 
            add_figure_to_doc(doc, fig, width_inches=w_full, height_inches=h_std)

        # -----------------------------------------------------------------
        # 3. Horizon Analysis Table
        # -----------------------------------------------------------------
        elif section_key == "horizon_table":
            check_break()
            add_header(doc, "Horizon Analysis", level=2)
            
            horizon_df = dw.get_horizon_analysis(data)
            if not horizon_df.empty:
                headers = ["Horizon", "Return", "P/L ($)"]
                # Convert DF to list of lists including index
                rows = []
                for idx, row in horizon_df.iterrows():
                    rows.append([
                        row.get('Horizon', str(idx)),
                        fmt_pct_clean(row.get('Return', 0)),
                        fmt_dollar_clean(row.get('P/L', 0))
                    ])
                add_table(doc, headers, rows, right_align=[1, 2])

        # -----------------------------------------------------------------
        # 4. Asset Allocation
        # -----------------------------------------------------------------
        elif section_key == "allocation":
            add_page_break(doc)
            add_header(doc, "Asset Allocation", level=2)
            
            # Unpack tuple
            res = dw.get_asset_allocation_charts(data, theme='light')
            if res:
                pie_fig = res[0] # Pie is first
                add_figure_to_doc(doc, pie_fig, width_inches=w_split, height_inches=h_std)
                
                check_break()

                bar_fig = res[1] 
                add_figure_to_doc(doc, bar_fig, width_inches=w_split, height_inches=h_std)

            # Allocation History Chart
            check_break()
            add_paragraph_centered(doc, "Allocation History", keep_after=True)
            hist_fig = dw.get_allocation_history_chart(data, theme='light')
            add_figure_to_doc(doc, hist_fig, width_inches=w_full, height_inches=h_std)

        # -----------------------------------------------------------------
        # 5. Sector Breakdown
        # -----------------------------------------------------------------
        elif section_key == "sector":
            check_break()
            add_header(doc, "Sector Breakdown", level=2)
            sector_fig = dw.get_sector_allocation_chart(data, theme='light')
            add_figure_to_doc(doc, sector_fig, width_inches=w_full, height_inches=h_std)

        # -----------------------------------------------------------------
        # 6. Top Holdings
        # -----------------------------------------------------------------
        elif section_key == "holdings":
            check_break()
            add_header(doc, "Top Holdings", level=2)
            sec_table = data.get('sec_table_current', pd.DataFrame()).copy()
            if not sec_table.empty:
                # Filter out CASH row
                if 'asset_class' in sec_table.columns:
                     sec_table = sec_table[sec_table['asset_class'] != 'CASH']
                     
                # Sort by weight if possible, else market value
                if 'weight' in sec_table.columns:
                    sort_col = 'weight'
                else:
                    sort_col = 'market_value'
                    
                sec_table = sec_table.sort_values(sort_col, ascending=False).head(10)
                
                headers = ["Ticker", "Asset Class", "Shares", "Value", "Weight"]
                rows = []
                for _, row in sec_table.iterrows():
                    rows.append([
                        str(row.get('ticker','')),
                        str(row.get('asset_class','')),
                        f"{row.get('shares', 0):,.2f}",
                        fmt_dollar_clean(row.get('market_value',0)),
                        fmt_pct_clean(row.get('weight',0))
                    ])
                    
                add_table(doc, headers, rows, right_align=[2, 3, 4])

        # -----------------------------------------------------------------
        # 7. Risk Metrics
        # -----------------------------------------------------------------
        elif section_key == "risk":
            check_break()
            add_header(doc, "Risk Metrics", level=2)
            risk_df = dw.get_risk_diversification(data) 
            if not risk_df.empty:
                 headers = ["Metric", "Value"]
                 rows = []
                 for _, row in risk_df.iterrows():
                     rows.append([
                         str(row.get('Metric', '')),
                         str(row.get('Value', ''))
                     ])
                 add_table(doc, headers, rows, right_align=[1])

        # -----------------------------------------------------------------
        # 8. Net Flows
        # -----------------------------------------------------------------
        elif section_key == "flows":
            check_break()
            add_header(doc, "Net Flows", level=2)
            
            # REPLACED: Use dynamic calculations based on report date range (not just YTD)
            # Replicates logic from dash_wrappers.get_flows_summary_ytd but with dynamic dates
            
            pv = data.get("pv")
            cf_ext = data.get("cf_ext")
            tx_raw = data.get("tx_raw")
            dividends = data.get("dividends")
            
            rows = []
            
            if not pv.empty:
                # Determine effective calculation window
                # Start: use start_date or inception
                eff_start = pd.Timestamp(start_date) if start_date else data.get("inception_date")
                # End: use end_date or max available
                eff_end = pd.Timestamp(end_date) if end_date else pv.index.max()
                
                # External Flows
                if cf_ext is not None and not cf_ext.empty:
                    flows_ext = cf_ext[(cf_ext["date"] >= eff_start) & (cf_ext["date"] <= eff_end)]
                    deposits = flows_ext.loc[flows_ext["amount"] > 0, "amount"].sum()
                    withdrawals = flows_ext.loc[flows_ext["amount"] < 0, "amount"].sum()
                    net_ext = flows_ext["amount"].sum()
                    most_recent_ext = flows_ext["date"].max()
                else:
                    deposits = 0.0
                    withdrawals = 0.0
                    net_ext = 0.0
                    most_recent_ext = pd.NaT

                # Internal Activity
                if tx_raw is not None and not tx_raw.empty:
                    tx_window = tx_raw[(tx_raw["date"] >= eff_start) & (tx_raw["date"] <= eff_end)]
                    buys = tx_window.loc[tx_window["amount"] < 0, "amount"].sum()
                    sells = tx_window.loc[tx_window["amount"] > 0, "amount"].sum()
                    most_recent_tx = tx_window["date"].max()
                else:
                    buys = 0.0
                    sells = 0.0
                    most_recent_tx = pd.NaT

                # Dividends
                if dividends is not None and not dividends.empty:
                    div_window = dividends[(dividends["date"] >= eff_start) & (dividends["date"] <= eff_end)]
                    income = div_window["amount"].sum()
                    most_recent_div = div_window["date"].max()
                else:
                    income = 0.0
                    most_recent_div = pd.NaT
                
                net_internal = buys + sells + income
                
                # Determine most recent date
                dates = [d for d in [most_recent_ext, most_recent_tx, most_recent_div] if pd.notna(d)]
                most_recent_str = max(dates).strftime("%Y-%m-%d") if dates else "N/A"
                
                # Format Rows
                rows = [
                    ["Net External Flows", fmt_dollar_clean(net_ext)],
                    ["• Deposits", fmt_dollar_clean(deposits)],
                    ["• Withdrawals", fmt_dollar_clean(withdrawals)],
                    ["Net Internal Activity", fmt_dollar_clean(net_internal)],
                    ["• Buys (Cash Out)", fmt_dollar_clean(buys)],
                    ["• Sells (Cash In)", fmt_dollar_clean(sells)],
                    ["• Income (Divs)", fmt_dollar_clean(income)],
                    ["Most Recent Flow", most_recent_str]
                ]

            if rows:
                 headers = ["Metric", "Value"]
                 add_table(doc, headers, rows, right_align=[1])

        # -----------------------------------------------------------------
        # 9. Tax Lot Explorer (Open Lots)
        # -----------------------------------------------------------------
        elif section_key == "tax_lots":
            check_break()
            title_text = "Tax Lot Explorer (Open Lots)"
            if end_date:
                title_text += f"\n(As of {pd.Timestamp(end_date).strftime('%Y-%m-%d')})"
            add_header(doc, title_text, level=2)
            
            # Use simple FIFO strat since we don't have signal easy access
            # NOW SUPPORTS TIME MACHINE VIA end_date
            calc_date = end_date if end_date else pd.Timestamp.now()
            open_lots, _ = build_tax_lots(strategy="FIFO", as_of_date=calc_date)
            
            if not open_lots.empty:
                # Ensure Date formatting
                if "Date Acquired" in open_lots.columns:
                    open_lots["Date Acquired"] = pd.to_datetime(open_lots["Date Acquired"]).dt.strftime("%Y-%m-%d")

                # Filter columns to match UI (hiding technical ones)
                cols_hide = []
                display_cols = [c for c in open_lots.columns if c not in cols_hide]
                
                # Format rows
                rows = []
                # Max width constraint - limit columns if too many
                # For Word, keep core columns
                core_cols = ["Ticker", "Date Acquired", "Shares", "Cost Basis", "Market Value", "Unrealized P/L", "Term"]
                # Fallback if names differ
                final_cols = [c for c in core_cols if c in display_cols]
                if not final_cols: final_cols = display_cols[:7]
                
                for _, row in open_lots[final_cols].iterrows():
                    r_vals = []
                    for col in final_cols:
                        val = row[col]
                        # Apply formatting based on column name
                        if col in ["Cost Basis", "Current Price", "Market Value", "Unrealized P/L", "Est Tax Liability"]:
                            r_vals.append(fmt_dollar_clean(val))
                        elif col == "Shares":
                            r_vals.append(f"{val:,.2f}")
                        else:
                            r_vals.append(str(val))
                    rows.append(r_vals)
                    
                # Assume last few columns are numeric and should be right-aligned
                add_table(doc, final_cols, rows)
            else:
                 doc.add_paragraph("No open tax lots found.")

        # -----------------------------------------------------------------
        # 10. Risk Analysis Charts
        # -----------------------------------------------------------------
        elif section_key == "risk_charts":
            add_page_break(doc)
            add_header(doc, "Risk Analysis Charts", level=2)
            
            # Risk Return
            add_paragraph_centered(doc, "Risk vs Return Profile", keep_after=True)
            fig_risk = dw.get_risk_return_chart(data, theme='light')
            add_figure_to_doc(doc, fig_risk, height_inches=h_std)
            
            check_break()

            # Correlation
            add_paragraph_centered(doc, "Correlation Matrix", keep_after=True)
            fig_corr = dw.get_correlation_heatmap(data, theme='light')
            add_figure_to_doc(doc, fig_corr, height_inches=h_std)
            
            check_break()
            
            # Drawdown
            if not mobile_mode: add_page_break(doc)
            add_paragraph_centered(doc, "Drawdown Analysis", keep_after=True)
            fig_dd = dw.get_drawdown_chart(data, theme='light')
            add_figure_to_doc(doc, fig_dd, height_inches=h_std)

        # -----------------------------------------------------------------
        # 11. Performance Deep Dive
        # -----------------------------------------------------------------
        elif section_key == "perf_deep_dive":
            add_page_break(doc)
            add_header(doc, "Performance Deep Dive", level=2)
            
            bm_map = {
                "S&P 500": "SPY",
                "Total US Market": "VTI",
                "Aggressive Alloc": "AOA"
            }
            add_paragraph_centered(doc, "Cumulative Return vs Benchmark", keep_after=True)
            cum_fig = dw.get_cumulative_return_chart(data, None, bm_map, theme='light')
            add_figure_to_doc(doc, cum_fig, height_inches=h_std)
            
            check_break()

            add_paragraph_centered(doc, "Growth of Invested Capital", keep_after=True)
            growth_fig = dw.get_growth_of_capital_chart(data, "Total", theme='light')
            add_figure_to_doc(doc, growth_fig, height_inches=h_std)

        # -----------------------------------------------------------------
        # 12. Attribution Analysis
        # -----------------------------------------------------------------
        elif section_key == "attribution":
            check_break()
            add_header(doc, "Attribution Analysis", level=2)
            attr_fig = dw.get_smart_attribution_chart(data, theme='light')
            add_figure_to_doc(doc, attr_fig, height_inches=5)

        # -----------------------------------------------------------------
        # 13. Asset Class Performance Table
        # -----------------------------------------------------------------
        elif section_key == "ac_perf_table":
            add_page_break(doc)
            add_header(doc, "Asset Class Performance", level=2)
            
            class_df = data.get('class_df', pd.DataFrame())
            sec_table = data.get('sec_table_current', pd.DataFrame())
            # Full list of horizons to match App
            horizons = ["1D", "1W", "MTD", "1M", "3M", "6M", "YTD", "1Y", "SI"]
            
            headers = ["Asset"] + horizons
            rows = []
            
            # Unique classes from class_df
            unique_classes = class_df['asset_class'].unique() if not class_df.empty else []
            
            for ac in unique_classes:
                # Map to Short Name if available
                ac_short = ASSET_CLASS_SHORT_MAP.get(ac, ac)
                
                # 1. Class Row
                crow_df = class_df[class_df['asset_class'] == ac]
                if crow_df.empty: continue
                crow = crow_df.iloc[0]
                
                r_vals = [ac_short]
                for h in horizons:
                    val = crow.get(h)
                    r_vals.append(fmt_pct_clean(val))
                rows.append(r_vals)
                
                # 2. Ticker Rows
                if not sec_table.empty:
                    tickers = sec_table[sec_table['asset_class'] == ac]
                    for _, trow in tickers.iterrows():
                        t = trow.get('ticker', '')
                        tr_vals = [f"    {t}"] # Indent
                        for h in horizons:
                            val = trow.get(h)
                            tr_vals.append(fmt_pct_clean(val))
                        rows.append(tr_vals)
                        
            add_table(doc, headers, rows, right_align=list(range(1, len(headers))))

        # -----------------------------------------------------------------
        # 14. Asset Class P/L Table
        # -----------------------------------------------------------------
        elif section_key == "ac_pl_table":
            add_header(doc, "Asset Class P/L (Economic)", level=2)
            
            class_df = data.get('class_df', pd.DataFrame())
            sec_table = data.get('sec_table_current', pd.DataFrame())
            horizons = ["1D", "1W", "MTD", "1M", "3M", "6M", "YTD", "1Y", "SI"]
            
            # Pre-fetch ticker P/L for all horizons to avoid re-calc in loop
            ticker_pl_cache = {}
            for h in horizons:
                try:
                    ticker_pl_cache[h] = dw.get_ticker_pl_df(data, h)
                except:
                    ticker_pl_cache[h] = pd.DataFrame() # Fallback
                
            rows = []
            headers = ["Asset"] + horizons
            
            # Determine order of asset classes (can rely on class_df order)
            unique_classes = class_df['asset_class'].unique() if not class_df.empty else []
            
            for ac in unique_classes:
                # Map to Short Name if available
                ac_short = ASSET_CLASS_SHORT_MAP.get(ac, ac)
                
                # 1. Class Row
                r_vals = [ac_short]
                for h in horizons:
                    res = dw.get_asset_class_pl(data, ac, h, return_components=False)
                    r_vals.append(fmt_dollar_clean(res))
                rows.append(r_vals)
                
                # 2. Ticker Rows belonging to this class
                if not sec_table.empty:
                    tickers = sec_table[sec_table['asset_class'] == ac]
                    for _, trow in tickers.iterrows():
                        t = trow['ticker']
                        tr_vals = [f"    {t}"] # Indent
                        for h in horizons:
                            pl_df = ticker_pl_cache.get(h)
                            val = None
                            if pl_df is not None and not pl_df.empty and t in pl_df.index:
                                val = pl_df.loc[t, "pl"]
                            tr_vals.append(fmt_dollar_clean(val))
                        rows.append(tr_vals)
                        
            add_table(doc, headers, rows, right_align=list(range(1, len(headers))))

    return doc
